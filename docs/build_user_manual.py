"""Generate the WowWed User Manual as a Word document."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Inches, Pt, RGBColor, Emu, Twips

# Brand colours (from WowWed UI)
TERRACOTTA = RGBColor(0xC9, 0x6A, 0x5A)
TERRACOTTA_DK = RGBColor(0x8B, 0x3A, 0x2E)
SAGE = RGBColor(0x4A, 0x7A, 0x56)
GOLD = RGBColor(0xB8, 0x86, 0x2E)
WARM_BROWN = RGBColor(0x4A, 0x35, 0x2C)
MUTED = RGBColor(0x6B, 0x55, 0x4A)
CREAM_HEX = "FDF6F2"
TERRACOTTA_HEX = "C96A5A"
SAGE_HEX = "6B9E78"
GOLD_HEX = "D4A84B"
SOFT_HEX = "FAF6F2"
TIP_HEX = "EEF6F0"
NOTE_HEX = "FFF8EE"
WARN_HEX = "FDF0EE"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="E8D4C8", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def prevent_table_split(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_table_width(table, width_cm):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="C96A5A", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    set_run_font(run, size=9, color=MUTED)


def add_num_pages(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    set_run_font(run, size=9, color=MUTED)


def set_section_header_footer(section, first_page=False):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("WowWed")
    set_run_font(r, size=10, bold=True, color=TERRACOTTA)
    r2 = hp.add_run("  ·  User Manual")
    set_run_font(r2, size=10, color=MUTED)
    add_bottom_border(hp, "E8D4C8", "8")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Confidential to WowWed users  ·  Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(fp)
    r = fp.add_run(" of ")
    set_run_font(r, size=9, color=MUTED)
    add_num_pages(fp)

    if first_page:
        section.different_first_page_header_footer = True
        fh = section.first_page_header
        fh.is_linked_to_previous = False
        fh.paragraphs[0].clear()
        ff = section.first_page_footer
        ff.is_linked_to_previous = False
        ffp = ff.paragraphs[0]
        ffp.clear()
        ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ffp.add_run("WowWed  ·  Plan Your Dream Wedding, Smarter.")
        set_run_font(r, size=9, italic=True, color=MUTED)


def para(doc, text, size=11, bold=False, italic=False, color=WARM_BROWN,
         align="left", space_before=0, space_after=8, first_line=0):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if first_line:
        p.paragraph_format.first_line_indent = Cm(first_line)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def mixed_para(doc, parts, align="left", space_before=0, space_after=8):
    """parts: list of (text, kwargs)."""
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run_font(r, **kwargs)
    return p


def heading1(doc, text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    label = f"{number}  {text}" if number else text
    r = p.add_run(label)
    set_run_font(r, name="Calibri", size=22, bold=True, color=TERRACOTTA_DK)
    add_bottom_border(p, TERRACOTTA_HEX, "18")
    return p


def heading2(doc, text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    label = f"{number}  {text}" if number else text
    r = p.add_run(label)
    set_run_font(r, size=15, bold=True, color=TERRACOTTA)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color=WARM_BROWN)
    return p


def bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.6)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, bold=True, color=WARM_BROWN)
        r = p.add_run(text)
        set_run_font(r, size=11, color=WARM_BROWN)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=WARM_BROWN)
    return p


def numbered(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, bold=True, color=WARM_BROWN)
        r = p.add_run(text)
        set_run_font(r, size=11, color=WARM_BROWN)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=WARM_BROWN)
    return p


def callout(doc, title, body, fill=TIP_HEX, accent=SAGE_HEX):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 16.5)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, accent, "12")
    set_cell_margins(cell, 80, 80, 120, 120)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=SAGE if fill == TIP_HEX else TERRACOTTA_DK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(body)
    set_run_font(r, size=10.5, color=WARM_BROWN)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def make_table(doc, headers, rows, col_widths=None, header_fill=TERRACOTTA_HEX):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 16.5)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_borders(cell, "B85A4A", "4")
        set_cell_margins(cell, 50, 50, 70, 70)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=RGBColor(255, 255, 255))

    for ri, row in enumerate(rows):
        fill = "FFFFFF" if ri % 2 == 0 else SOFT_HEX
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            shade_cell(cell, fill)
            set_cell_borders(cell, "E8D4C8", "4")
            set_cell_margins(cell, 50, 50, 70, 70)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            set_run_font(r, size=10, color=WARM_BROWN)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    set_section_header_footer(section, first_page=True)

    # Styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = WARM_BROWN

    # =====================================================================
    # COVER
    # =====================================================================
    for _ in range(3):
        para(doc, "", space_after=0)

    p = para(doc, "WOWWED", size=14, bold=True, color=TERRACOTTA, align="center", space_after=4)
    p = para(doc, "User Manual", size=36, bold=True, color=TERRACOTTA_DK, align="center", space_after=6)
    para(
        doc,
        "A complete guide for couples and wedding vendors",
        size=14,
        italic=True,
        color=MUTED,
        align="center",
        space_after=18,
    )

    # Accent bar via 1-col table
    bar = doc.add_table(rows=1, cols=1)
    set_table_width(bar, 10)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = bar.cell(0, 0)
    shade_cell(c, TERRACOTTA_HEX)
    set_cell_borders(c, TERRACOTTA_HEX, "0")
    c.paragraphs[0].clear()
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    r = c.paragraphs[0].add_run(" ")
    set_run_font(r, size=4)

    para(doc, "", space_after=14)
    para(
        doc,
        "Plan Your Dream Wedding, Smarter.",
        size=16,
        italic=True,
        color=TERRACOTTA,
        align="center",
        space_after=20,
    )

    cover_meta = [
        ("Product", "WowWed — wedding planning platform"),
        ("Audience", "Couples and vendors in Sri Lanka"),
        ("Version", "1.0"),
        ("Date", date.today().strftime("%d %B %Y")),
        ("Language", "English"),
    ]
    meta = doc.add_table(rows=len(cover_meta), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, 11)
    for i, (k, v) in enumerate(cover_meta):
        left, right = meta.rows[i].cells
        shade_cell(left, SOFT_HEX)
        shade_cell(right, "FFFFFF")
        set_cell_borders(left, "E8D4C8", "4")
        set_cell_borders(right, "E8D4C8", "4")
        set_cell_margins(left, 50, 50, 90, 90)
        set_cell_margins(right, 50, 50, 90, 90)
        left.paragraphs[0].clear()
        r = left.paragraphs[0].add_run(k)
        set_run_font(r, size=11, bold=True, color=TERRACOTTA_DK)
        right.paragraphs[0].clear()
        r = right.paragraphs[0].add_run(v)
        set_run_font(r, size=11, color=WARM_BROWN)

    para(doc, "", space_after=28)
    para(
        doc,
        "This manual describes how to create an account, plan a wedding, manage guests and budget,\n"
        "book vendors, design invitations, and use WowWed’s smart planning tools.",
        size=10.5,
        color=MUTED,
        align="center",
        space_after=6,
    )
    para(
        doc,
        "© WowWed. For use by registered couples and vendors.",
        size=9,
        color=MUTED,
        align="center",
        space_after=0,
    )

    page_break(doc)

    # =====================================================================
    # CONTENTS
    # =====================================================================
    heading1(doc, "Contents")
    para(
        doc,
        "Use this list to jump to the chapter you need. Button names and menu labels in this manual match the wording on screen.",
        align="justify",
    )

    toc = [
        ("1", "Welcome to WowWed", "3"),
        ("2", "Getting started", "4"),
        ("3", "Your couple dashboard", "6"),
        ("4", "Wedding checklist", "7"),
        ("5", "Guest list", "8"),
        ("6", "Seating chart", "9"),
        ("7", "Budget and expenses", "10"),
        ("8", "Vendors and booking requests", "12"),
        ("9", "Invitations", "14"),
        ("10", "Analytics, reports and crew", "15"),
        ("11", "WowBot assistant", "16"),
        ("12", "Account settings", "17"),
        ("13", "Vendor portal", "18"),
        ("14", "Smart planning tools", "20"),
        ("15", "Tips, troubleshooting and glossary", "21"),
    ]
    for num, title, _pg in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{num}    {title}")
        set_run_font(r, size=12, color=WARM_BROWN)

    page_break(doc)

    # =====================================================================
    # 1 WELCOME
    # =====================================================================
    heading1(doc, "Welcome to WowWed", "1")
    heading2(doc, "What WowWed is", "1.1")
    para(
        doc,
        "WowWed is a free wedding planning platform built for Sri Lankan weddings. "
        "It brings your checklist, guest list, seating chart, budget, vendor bookings and invitations "
        "into one dashboard so you can plan a Poruwa, Church, Hindu Tamil, Muslim Nikah or reception wedding "
        "across all 25 districts.",
        align="justify",
    )
    para(
        doc,
        "The public site tagline is “Plan Your Dream Wedding, Smarter.” Couples do not need a credit card. "
        "Amounts throughout the app are in Sri Lankan Rupees (LKR / Rs.).",
        align="justify",
    )

    heading2(doc, "Who this manual is for", "1.2")
    make_table(
        doc,
        ["You are", "You will use WowWed to"],
        [
            ["A couple", "Create a wedding profile, track tasks and guests, plan seating and budget, book vendors, and design invitations."],
            ["A vendor", "Publish a live listing, receive booking requests, reply or negotiate, and manage your calendar."],
            ["A visitor", "Read about WowWed on the home page, start planning, or chat with WowBot before you create an account."],
        ],
        col_widths=[4, 12.5],
    )
    para(
        doc,
        "Wedding guests do not log in. You record them on the Guest List. There is no separate admin or wedding-planner login. "
        "Crew members (bridesmaids, best man, and so on) are names you assign to tasks — they do not have their own accounts.",
        align="justify",
    )

    heading2(doc, "What you can do", "1.3")
    bullet(doc, "Sign up as a couple or as a vendor in three short steps.")
    bullet(doc, "See a readiness score on your home page (checklist, RSVPs and budget).")
    bullet(doc, "Work through a Sri Lankan wedding checklist with 70+ starter tasks.")
    bullet(doc, "Import guests, track RSVPs, and export a CSV for Excel.")
    bullet(doc, "Auto-seat guests who marked Coming, then fine-tune tables by hand.")
    bullet(doc, "Estimate vendor cost, save budget scenarios, and log real expenses.")
    bullet(doc, "Search local vendors, send booking requests, and confirm a hire.")
    bullet(doc, "Design a printable invitation PDF from cultural templates.")
    bullet(doc, "Download guest, budget, vendor, seating and full-wedding PDF reports.")
    bullet(doc, "Ask WowBot planning questions — it replies from WowWed’s knowledge base, not a live chat agent.")

    heading2(doc, "How to read this manual", "1.4")
    para(
        doc,
        "Button and menu names appear in bold, exactly as they do in the app — for example Log in, Auto-seat all, or Send booking request. "
        "Tips sit in green boxes. Notes and cautions sit in cream or rose boxes.",
        align="justify",
    )
    callout(
        doc,
        "Tip",
        "Complete your wedding profile first. District, ceremony type, date and total budget drive Vendor Match, cost estimates, analytics and WowBot answers.",
    )

    # =====================================================================
    # 2 GETTING STARTED
    # =====================================================================
    heading1(doc, "Getting started", "2")
    heading2(doc, "Open WowWed", "2.1")
    para(
        doc,
        "Open the WowWed website in your browser. The home page shows About, Features and Vendors. "
        "From the navigation bar you can Start planning (couples), Join as a vendor, or Log in.",
        align="justify",
    )

    heading2(doc, "Create a couple account", "2.2")
    para(doc, "Couples complete three steps: a short setup, an account, then a wedding profile.", align="justify")
    heading3(doc, "Step 1 — Almost there!")
    para(doc, "Click Start planning free or Start planning. You land on the couple setup page.", align="justify")
    bullet(doc, " Wedding location (Optional) — choose a district.", bold_lead="Wedding location (Optional)")
    bullet(doc, " — optional.", bold_lead="Wedding Date (Optional)")
    bullet(doc, " — Indoor, Outdoor or Mixed. Required.", bold_lead="Venue Type *")
    bullet(doc, " — Poruwa Ceremony, Church Wedding, Hindu Tamil Wedding, Muslim Nikah Ceremony, or Reception. Required.", bold_lead="Ceremony Type *")
    para(doc, "Click Complete Setup. New visitors continue to Create your account. If you are already logged in, you go to the wedding profile.")

    heading3(doc, "Step 2 — Create your account")
    para(doc, "The subtitle is “Next you’ll create your wedding profile.” Fill in:", align="justify")
    make_table(
        doc,
        ["Field", "What to enter"],
        [
            ["Full name", "Your name (example placeholder: Priya Perera)."],
            ["Email", "The address you will use to log in."],
            ["Phone", "A number vendors can reach."],
            ["Password", "At least 6 characters."],
            ["Confirm password", "Type the same password again."],
        ],
        col_widths=[5, 11.5],
    )
    para(doc, "Click Create account. You move to the wedding profile with your new account signed in.")

    heading3(doc, "Step 3 — Create your wedding profile")
    para(doc, "This profile is the heart of WowWed. Fill in:", align="justify")
    make_table(
        doc,
        ["Field", "Required?", "Notes"],
        [
            ["Partner 1 name", "Yes", "First partner as you want it shown."],
            ["Partner 2 name", "Yes", "Second partner as you want it shown."],
            ["Wedding location", "Yes", "District — used for vendor matching."],
            ["Wedding date", "Yes", "Drives the countdown and checklist timeline."],
            ["Ceremony type", "Yes", "Poruwa, Church, Hindu Tamil, Nikah, or Reception."],
            ["Expected guests", "No", "Fallback headcount if RSVPs are not in yet."],
            ["Wedding scale", "No", "Budget · simple / Standard · typical / Premium · grand."],
            ["Total budget (LKR)", "Yes", "Your overall ceiling for estimates and alerts."],
        ],
        col_widths=[4.5, 2.5, 9.5],
    )
    para(doc, "Click Save and go to dashboard. WowWed seeds your checklist with Sri Lankan wedding tasks. During signup you may also choose ← Skip to dashboard if you want to finish the profile later.")

    heading2(doc, "Create a vendor account", "2.3")
    para(doc, "Click Join as a vendor. Vendors also complete three steps.", align="justify")
    heading3(doc, "Step 1 — Where are you as a vendor?")
    para(doc, "Choose the option that fits, then click Next:", align="justify")
    bullet(doc, "Just starting my wedding business")
    bullet(doc, "I have a business but I'm new to WowWed")
    bullet(doc, "Listed elsewhere and want to reach more couples")
    bullet(doc, "Established and ready to receive bookings")

    heading3(doc, "Step 2 — Almost there!")
    bullet(doc, " tap every category that applies: Venue & Res. Halls, Bridal Service, Groom service, Photography & Videography, Jewellary, Floral & Deco, Caters, Cakes.", bold_lead="Service categories —")
    bullet(doc, " use Search districts… and select every district you serve (all 25 Sri Lankan districts are listed).", bold_lead="Business districts —")
    para(doc, "Click Complete Setup, then create your account the same way as a couple. After Create account you go to My Listing to publish your public page.")

    heading2(doc, "Log in", "2.4")
    para(doc, "Open Log in. The page title is Welcome back, with the line “Log in to continue planning your wedding.”", align="justify")
    numbered(doc, "Enter your Email and Password.")
    numbered(doc, "Click Log in.")
    numbered(doc, "Couples land on the dashboard Home. Vendors land on the vendor Overview.")
    para(doc, "Links at the bottom: Forgot password? and Get started.")

    heading2(doc, "Reset your password", "2.5")
    numbered(doc, "Click Forgot password? on the login page.")
    numbered(doc, "On Reset password, enter your Email and a New password (at least 6 characters).")
    numbered(doc, "Click Update password. After a short success message you return to login.")

    callout(
        doc,
        "Note",
        "A couple account cannot open the vendor pages, and a vendor account cannot open the couple dashboard. Use the matching Get started path if you need both kinds of access.",
        fill=NOTE_HEX,
        accent=GOLD_HEX,
    )

    # =====================================================================
    # 3 DASHBOARD HOME
    # =====================================================================
    heading1(doc, "Your couple dashboard", "3")
    heading2(doc, "The sidebar", "3.1")
    para(doc, "After you log in as a couple, a sidebar stays on the left. Your names at the top open Edit wedding profile.", align="justify")
    make_table(
        doc,
        ["Sidebar label", "Opens"],
        [
            ["Home", "Wedding overview and readiness score"],
            ["Wedding Checklist", "Tasks and timeline"],
            ["Guest List", "Guests and RSVPs"],
            ["Seating Chart", "Tables and auto-seat"],
            ["Invitations", "Invitation studio"],
            ["Budget", "Cost estimate and tracking"],
            ["Vendors", "Browse and match vendors"],
            ["Requests", "Your booking requests"],
            ["Analytics", "Charts and wedding-at-a-glance"],
            ["Reports", "Downloadable PDF reports"],
            ["Assistant", "Full-page WowBot"],
            ["Wedding Crew", "Bridal party and helpers"],
            ["Settings", "Account and logout (bottom of sidebar)"],
        ],
        col_widths=[5, 11.5],
    )
    para(doc, "The public navbar and footer are hidden on the dashboard. A notification bell sits in the top bar.")

    heading2(doc, "Home", "3.2")
    para(
        doc,
        "Home greets you with Good morning, Good afternoon or Good evening and the line “your wedding home.” "
        "You see your names, a wedding-date chip, district, ceremony type, and a countdown such as “X days until the big day.”",
        align="justify",
    )
    heading3(doc, "Wedding readiness score")
    para(doc, "The score combines three parts of your plan:", align="justify")
    bullet(doc, "Checklist completion — 40%")
    bullet(doc, "Guest RSVPs — 35%")
    bullet(doc, "Budget use — 25%")
    para(doc, "A green On track badge appears at 70% or above. Yellow and red appear when the score is lower.")
    para(doc, "Glance cards show tasks completed, guests confirmed, and budget used. Your planning tools grid links to every module with live counts. Coming up next lists the next six incomplete checklist tasks.")

    heading2(doc, "Notifications", "3.3")
    para(doc, "Open the bell in the top bar. Local reminders you can dismiss include:", align="justify")
    bullet(doc, "Wedding countdown when the date is 30 days away or closer.")
    bullet(doc, "Guest RSVPs still waiting.")
    bullet(doc, "More than five checklist tasks still to do.")
    bullet(doc, "Budget alert or Budget nearly used when spend passes 90% of your total.")
    bullet(doc, "Vendor requests waiting for a reply, or Ready to confirm when a vendor has accepted.")
    para(doc, "Inbox messages from the server also appear here (for example a vendor reply). The list refreshes in the background while you work.")

    # =====================================================================
    # 4 CHECKLIST
    # =====================================================================
    heading1(doc, "Wedding checklist", "4")
    para(
        doc,
        "Open Wedding Checklist from the sidebar. WowWed seeds 70 or more Sri Lankan tasks when you first save your wedding profile "
        "(Poruwa items, Ashlaka / Jayamangala Gatha, and similar). Task phases follow your wedding date.",
        align="justify",
    )
    heading2(doc, "Read the summary", "4.1")
    para(doc, "The bar at the top shows how many tasks are done, how many are to do, percent complete, your wedding date, and how many tasks are unassigned.")

    heading2(doc, "Find a task", "4.2")
    bullet(doc, "Type in Search tasks by name…")
    bullet(doc, "Your timeline: All stages, Just started, 6 months to go, 3 months to go, 1 month to go, 1 week to go, Wedding day.")
    bullet(doc, "Status: All, To do, Done.")
    bullet(doc, "Category: Guests, Suite and Dress, Vendors, Ceremony, Catering, Decorations, Entertainment, Logistics, Budget, Venue, Other.")
    bullet(doc, "Assigned to: filter unassigned tasks, or filter by crew role.")

    heading2(doc, "Complete or edit a task", "4.3")
    numbered(doc, "Tick the checkbox on a row to mark the task complete.")
    numbered(doc, "Click the row or the edit (pencil) icon to open the panel.")
    numbered(doc, "Change Task name, Status (To do / Completed), Due date, Category, Assign to (Unassigned or a crew role), and Notes.")
    numbered(doc, "Click Save changes, or Cancel.")

    heading2(doc, "Add a task", "4.4")
    numbered(doc, "Click + Add task.")
    numbered(doc, "Enter a name and the same fields as above.")
    numbered(doc, "Save. Assign people first on Wedding Crew if you want names in Assign to.")
    callout(
        doc,
        "Tip",
        "Use the timeline filters as the date approaches. “Coming up next” on Home always pulls from this list, so keeping due dates honest keeps Home useful.",
    )

    # =====================================================================
    # 5 GUESTS
    # =====================================================================
    heading1(doc, "Guest list", "5")
    para(
        doc,
        "Open Guest List. This list feeds seating, budget guest counts, analytics and reports. "
        "RSVP labels in the app are Waiting, Coming and Not coming.",
        align="justify",
    )

    heading2(doc, "Add one guest", "5.1")
    numbered(doc, "Click + Add guest.")
    numbered(doc, "Enter Full name * (required). Optionally add Phone, Email, Group, RSVP status, Age (for example 8 or 72), and Notes.")
    numbered(doc, "If two people should not share a table, type names in Do not sit with (names, comma separated).")
    numbered(doc, "Click Add guest.")

    heading2(doc, "Guest groups", "5.2")
    para(doc, "Groups help filters, seating and reports:", align="justify")
    para(
        doc,
        "No Group · Bride's Family · Groom's Family · Bride's Friends · Groom's Friends · "
        "Bride's Colleagues · Groom's Colleagues · Relatives · VIP · Neighbours · Other",
        italic=True,
        color=MUTED,
    )

    heading2(doc, "Add many guests at once", "5.3")
    numbered(doc, "Open More → Add many guests.")
    numbered(doc, "Paste names, one per line. Choose a group and RSVP for the batch.")
    numbered(doc, "Click Add all guests.")

    heading2(doc, "Import and export CSV", "5.4")
    para(doc, "Open the More menu for Export CSV, Import / update CSV, and Delete imported CSV.", align="justify")
    para(doc, "Column order for a spreadsheet:", align="justify")
    para(doc, "name, email, phone, group, rsvp, age, notes, avoid", italic=True, color=MUTED)
    callout(
        doc,
        "Tip",
        "Export CSV, update RSVPs in Excel, then Import / update CSV to refresh many replies at once. A collapsible tip on the page describes this workflow.",
    )

    heading2(doc, "Filter and update RSVPs", "5.5")
    bullet(doc, "Use the chips All, Waiting, Coming, Not coming.")
    bullet(doc, "Search by name, email, or phone. Narrow by Group or RSVP, then Clear.")
    bullet(doc, "Select rows, then Mark coming, Mark not coming, Mark waiting, Move to group…, or Delete selected.")
    para(doc, "Open a guest to Save changes or Delete guest.")

    # =====================================================================
    # 6 SEATING
    # =====================================================================
    heading1(doc, "Seating chart", "6")
    para(
        doc,
        "Open Seating Chart. The summary shows how many tables you have, how many guests are seated, and either “N need a chair” or All guests seated.",
        align="justify",
    )

    heading2(doc, "Add and style tables", "6.1")
    numbered(doc, "Click + Add table. Set name, seats, shape, suite, priority and a preferred guest group.")
    numbered(doc, "Or open Table settings (shape & chairs for all tables) and apply Round, Rectangle, Square, Head Table or Standing, then adjust Chairs per table with − / +.")
    para(doc, "Suite zones (visual grouping): General, VIP, Bride Family, Groom Family, Friends.")
    para(doc, "Tap a table to open its roster. Use Seat with group, Clear this table, edit, or delete.")

    heading2(doc, "Seat people by hand", "6.2")
    numbered(doc, "Open the Guests tab in the sidebar. Search guests… or filter by Guest group.")
    numbered(doc, "Switch between seated and waiting to seat.")
    numbered(doc, "Click a guest, then click an empty chair.")

    heading2(doc, "Auto-seat all (smart seating)", "6.3")
    numbered(doc, "On the Guest List, mark people who will attend as Coming. Auto-seat only uses Coming guests.")
    numbered(doc, "Add tables, or let WowWed add tables if there are not enough chairs.")
    numbered(doc, "Click Auto-seat all.")
    numbered(doc, "Read any Review after Auto-seat messages. These are suggestions, not errors. Move people if you disagree.")
    callout(
        doc,
        "How smart seating works",
        "WowWed tries a clustering model with guest-group rules and your Do not sit with names. If that service is offline, a built-in rule-based layout is used instead. Either way you can still drag people by clicking a guest and an empty chair.",
    )

    # =====================================================================
    # 7 BUDGET
    # =====================================================================
    heading1(doc, "Budget and expenses", "7")
    para(
        doc,
        "Open Budget. The page has a planning form (“Fill in one form, get a cost estimate, save plans, and pick your favourite.”), "
        "live tracking against Total budget (LKR) from your profile, and hired-vendor payments.",
        align="justify",
    )

    heading2(doc, "Plan your wedding budget", "7.1")
    para(doc, "Adjust the form. Guests coming (RSVP) updates from the Guest List automatically. Typical fields include:", align="justify")
    bullet(doc, "Wedding style (Budget / Standard / Premium scale), District, Ceremony, Venue, Season (Regular or Peak +8%).")
    bullet(doc, "Meals & catering, Plate price, Drinks, Reception time.")
    bullet(doc, "Photography, Videography, Bridal outfit, Groom outfit, Hair & makeup.")
    bullet(doc, "Décor & flowers, Lighting, Entertainment, Wedding cake, Jewellery, Invitations, Transport.")
    para(doc, "Most line items offer Custom — add your own with your own amounts.")
    para(doc, "Results show estimated vendor cost, a low–high range, a per-guest figure, and a verdict against Your budget. A Cost breakdown lists Venue, Catering, Photography and the rest.")
    bullet(doc, "Save scenario — keep a named plan you can load later.")
    bullet(doc, "Confirm plan — lock the plan you prefer.")
    bullet(doc, "Reset to wedding profile — restore values from your profile.")
    bullet(doc, "Try again — if the estimate service fails.")

    heading2(doc, "Track spend", "7.2")
    para(doc, "You see total budget, spent, remaining, and an overspend alert if you go past the ceiling.", align="justify")
    para(doc, "Default categories: Venue, Catering, Photography, Videography, Attire, Decorations, Entertainment, Transport, Beauty, Ceremony, Invitations, Vendors.")
    numbered(doc, "Click + Add category to create another envelope, or edit allocated amounts.")
    numbered(doc, "Sort by Default, Name, Remaining budget or Budget allocated.")
    numbered(doc, "Click + Add expense — name, amount, category, date, notes.")
    numbered(doc, "Click View all expenses to open Your expense overview.")

    heading2(doc, "Your expense overview", "7.3")
    para(doc, "This sub-page lists Name, Category, Date and Amount. Filter All, Uncategorized or a single category. Use + Add new expense or ← Back to budget. The footer shows Total expenses to date.")

    heading2(doc, "Hired vendors on the budget", "7.4")
    para(
        doc,
        "When you Confirm booking on Requests, WowWed counts that hire toward spend (shown as Paid). "
        "You do not enter a card on WowWed — “Paid” means you confirmed the booking in the planner, not that a payment gateway ran.",
        align="justify",
    )
    callout(
        doc,
        "Tip",
        "Keep Expected guests on the profile as a fallback, but mark real RSVPs as Coming. The budget form’s guest count follows confirmed guests when they exist.",
    )

    # =====================================================================
    # 8 VENDORS (COUPLE)
    # =====================================================================
    heading1(doc, "Vendors and booking requests", "8")
    heading2(doc, "Find vendors", "8.1")
    para(doc, "Open Vendors. Use Category (All Categories plus the eight vendor types), Location (All districts), and Search vendors, places…", align="justify")
    para(doc, "Turn on Match for Vendor Match. Ranking uses your district, budget and ceremony type from the wedding profile. Open More filters for Price & rating filters (price bands; Rating Any / 4.0+ / 4.5+ / 4.8+).")
    para(doc, "On a card click View details (or View request if you already wrote to them). Select two or three listings and click Compare vendors.")
    para(doc, "Your chosen vendors strip links to Manage requests / Open requests.")

    heading2(doc, "Send a booking request", "8.2")
    numbered(doc, "Open a vendor and click Send booking request.")
    numbered(doc, "Set the event date, Your budget for this vendor, and a short note in Tell them about your wedding plans…")
    numbered(doc, "Submit. WowWed checks the vendor calendar for that date.")
    numbered(doc, "Track the request under Requests in the sidebar.")

    heading2(doc, "Manage requests", "8.3")
    para(doc, "The page title is Vendor requests. Tabs: Need action, Waiting, Booked, Other, All.", align="justify")
    make_table(
        doc,
        ["When this happens", "You can"],
        [
            ["You sent a request", "Wait on Waiting. Cancel if you change your mind."],
            ["Vendor accepted", "Need action — Confirm booking (adds the amount as Paid on Budget) or Negotiate reply."],
            ["Vendor sent an offer", "Accept this offer or Negotiate reply."],
            ["Hire is confirmed", "Booked — View in budget. Leave a star rating (1–5), write a Review, Submit review."],
            ["Rejected or cancelled", "Other."],
        ],
        col_widths=[5.5, 11],
    )
    para(doc, "If the list is empty, use Find vendors.")

    heading2(doc, "What each status means", "8.4")
    make_table(
        doc,
        ["Couple tab", "Vendor tab", "Meaning"],
        [
            ["Waiting", "Needs reply", "Request sent; vendor has not decided."],
            ["Need action", "Awaiting couple", "Vendor accepted — you confirm or negotiate."],
            ["Need action", "Needs reply", "A counter-offer is in play."],
            ["Booked", "Paid", "You confirmed the hire; budget updates."],
            ["Other", "—", "Rejected or cancelled."],
        ],
        col_widths=[4.5, 4.5, 7.5],
    )

    # =====================================================================
    # 9 INVITATIONS
    # =====================================================================
    heading1(doc, "Invitations", "9")
    para(doc, "Open Invitations. The studio has three steps: Pick design, Your details, Customize.", align="justify")

    heading2(doc, "Step 1 — Which style do you love?", "9.1")
    bullet(doc, "Search gold, nikkah, floral, blank…")
    bullet(doc, "Filter: All Templates, Sinhala, Church, Tamil, Muslim, Luxury, Modern.")
    bullet(doc, "Blank Canvas — “Design it yourself.”")
    para(doc, "Click Next — add your details →")

    heading2(doc, "Step 2 — Your details", "9.2")
    bullet(doc, "Click ✨ Fill from my profile to pull names and date.")
    bullet(doc, "💍 First partner's name, 💍 Second partner's name, 📅 Wedding date, 🕐 Time, 🏛️ Venue name.")
    bullet(doc, "Open ▸ Parents, address & RSVP for parents, full address and RSVP phone.")
    para(doc, "Click Design my invitation → or ← Back.")

    heading2(doc, "Step 3 — Customize and download", "9.3")
    para(doc, "Use the toolbar to add a heading, names, date, quote, text or photo. All moves or styles pieces together.", align="justify")
    para(doc, "Click ⬇ Download your invitation PDF. Use Change design or Edit details if you need to go back.")

    # =====================================================================
    # 10 ANALYTICS, REPORTS, CREW
    # =====================================================================
    heading1(doc, "Analytics, reports and crew", "10")
    heading2(doc, "Analytics", "10.1")
    para(doc, "Open Analytics (title: Analytics Dashboard). Your wedding at a glance lists date, district, ceremony, guests, budget, scale, checklist, RSVPs, spend, bookings, seating and crew.", align="justify")
    para(doc, "Charts include RSVP breakdown (Accepted / Pending / Rejected), task completion, budget utilisation, readiness score, and spending by category. After Auto-seat all you may also see Smart seating quality (accuracy, table fit, capacity notes). Edit profile opens the wedding profile.")

    heading2(doc, "PDF reports", "10.2")
    para(doc, "Open Reports (title: PDF Reports). Each card has Download PDF. The browser print window opens — choose Save as PDF.", align="justify")
    make_table(
        doc,
        ["Report", "What you get"],
        [
            ["Guest List Report", "Every guest with RSVP status and group."],
            ["Budget Report", "Planned versus actual by category."],
            ["Vendor Report", "Booking requests and vendor expenses."],
            ["Seating Chart", "Printable Find Your Seat chart."],
            ["Full Wedding Summary", "Complete overview in one file."],
        ],
        col_widths=[5, 11.5],
    )

    heading2(doc, "Wedding crew", "10.3")
    numbered(doc, "Open Wedding Crew and click + Add Member.")
    numbered(doc, "Enter Name * and a Role: Bridesmaid, Groomsman, Best Man, Maid of Honour, Flower Girl, Ring Bearer, Helper, Coordinator.")
    para(doc, "Those roles appear in Assign to on the checklist. Crew members do not receive a login.")

    # =====================================================================
    # 11 WOWBOT
    # =====================================================================
    heading1(doc, "WowBot assistant", "11")
    para(
        doc,
        "WowBot is WowWed’s keyword assistant. The launcher says Chat with WowBot. The header is Chat with WowBot 👋 with “We reply immediately.” "
        "The footer reads Powered by WowWed. It is not a live human and not a generative chat AI — answers come from WowWed’s built-in knowledge base.",
        align="justify",
    )
    para(doc, "The floating widget appears on most pages. It is hidden on vendor pages (so vendors are not shown couple-focused help) and on the dashboard Assistant page, where the same chat is shown full screen. Tagline on that page: Offline chatbot.")

    heading2(doc, "Ask a question", "11.1")
    numbered(doc, "Click Chat with WowBot, or open Assistant in the sidebar.")
    numbered(doc, "Pick a topic: Ideas, To-do list, Ceremony, Guests, Seating, Invites, Vendors, Money — or type in Ask about guests, budget, or vendors…")
    numbered(doc, "Click Send (or the send arrow). Use suggestion chips after a reply.")
    numbered(doc, "Click Start over (↺ on the widget) to clear the thread.")

    heading2(doc, "Starter questions", "11.2")
    bullet(doc, "What can you do?")
    bullet(doc, "What is WOWWED?")
    bullet(doc, "Tell me about our wedding")
    bullet(doc, "Help me plan my wedding")
    bullet(doc, "How can we arrange the seating?")
    bullet(doc, "How should we split the budget?")

    heading2(doc, "Personal answers and budget splits", "11.3")
    para(
        doc,
        "When you are logged in as a couple, WowBot can read your profile and live numbers — names, date, district, RSVPs, spend, checklist progress, table count and bookings. "
        "Questions such as “How many guests confirmed?” can include real counts and Open Guest List or Open Budget links.",
        align="justify",
    )
    para(
        doc,
        "If you write a total such as “Our budget is Rs. 2 million,” WowBot can suggest a split: Venue 35%, Catering 22%, Photo 12%, Attire 10%, Décor 8%, Extras 8%, Buffer 5%. Treat this as a starting point, then refine it on Budget.",
        align="justify",
    )

    # =====================================================================
    # 12 SETTINGS
    # =====================================================================
    heading1(doc, "Account settings", "12")
    para(doc, "Open Settings. Tagline: Your account and wedding details.", align="justify")
    heading3(doc, "Account card")
    para(doc, "Shows Name, Email, Phone and Role. Use Reset password or Log out.")
    heading3(doc, "Wedding card")
    para(doc, "Shows a profile summary and planning stage. Click Edit wedding profile to change names, date, district, ceremony, guests, scale and Total budget (LKR).")
    callout(
        doc,
        "Note",
        "Changing district, ceremony or budget updates Vendor Match, cost estimates and analytics. Changing the wedding date updates the countdown and checklist timeline.",
        fill=NOTE_HEX,
        accent=GOLD_HEX,
    )

    # =====================================================================
    # 13 VENDOR PORTAL
    # =====================================================================
    heading1(doc, "Vendor portal", "13")
    para(
        doc,
        "Vendors use three pages: Overview, Bookings and My Listing. There is no couple-style WowBot on these routes. "
        "A listing goes live as soon as you click Publish listing — there is no approval wait.",
        align="justify",
    )

    heading2(doc, "Publish your listing", "13.1")
    para(doc, "Open My Listing (title: Your public listing). Fill in Business details:", align="justify")
    bullet(doc, "Business name *")
    bullet(doc, "Categories (the same eight types couples search)")
    bullet(doc, "Districts served")
    bullet(doc, "Price range — Budget, Standard, Premium, Luxury, or a custom range")
    bullet(doc, "Description")
    para(doc, "Portfolio photos: up to 6 images, each 900 KB or smaller. Service packages / quotations: title, price, details, and an optional PDF up to 2 MB.")
    para(doc, "Preview shows How couples will see you, including a Send booking request button. Click Publish listing the first time, or Save changes later. Listing strength on Overview is a checklist of these fields — Complete listing or Edit listing jumps back here.")

    heading2(doc, "Overview", "13.2")
    para(doc, "The hero shows your business name, categories, districts, star rating and a Live pill. Stats: Needs reply, Awaiting couple, Paid, Earnings.", align="justify")
    para(doc, "Incoming requests appear here with Accept, Reply, Reject and Negotiate. You also see Next step, Upcoming dates, Earnings, Reviews, and listing strength. Buttons: Reply to N / Bookings, Edit listing.")

    heading2(doc, "Booking requests and calendar", "13.3")
    para(doc, "Open Bookings (title: Booking requests). Tabs: Needs reply, Awaiting couple, Paid, All. Download a booking summary or Download PDF summary.", align="justify")
    heading3(doc, "Reply to a couple")
    bullet(doc, " — status becomes Confirmed; the couple must still confirm.", bold_lead="Accept")
    bullet(doc, " — send a message.", bold_lead="Reply")
    bullet(doc, " — close the request.", bold_lead="Reject")
    bullet(doc, " / Hide offer — send a counter-offer with amount and message, then Send offer.", bold_lead="Negotiate")
    heading3(doc, "Calendar")
    para(doc, "The month calendar at the bottom uses Today and a legend: Pending, Negotiate, Confirmed, Paid. Click a date to filter the list; Clear removes the date filter.")

    heading2(doc, "Vendor categories", "13.4")
    make_table(
        doc,
        ["Category on WowWed", "Typical services"],
        [
            ["Venue & Res. Halls", "Halls, hotels, outdoor grounds"],
            ["Bridal Service", "Bridal wear and styling"],
            ["Groom service", "Groom wear and grooming"],
            ["Photography & Videography", "Photo and film packages"],
            ["Jewellary", "Bridal and couple jewellery"],
            ["Floral & Deco", "Poruwa, stage and floral décor"],
            ["Caters", "Menus and service staff"],
            ["Cakes", "Wedding cakes"],
        ],
        col_widths=[6, 10.5],
    )

    # =====================================================================
    # 14 SMART TOOLS
    # =====================================================================
    heading1(doc, "Smart planning tools", "14")
    para(doc, "WowWed labels these as Decision support on the home page. You do not need to understand the models to use the buttons.", align="justify")
    make_table(
        doc,
        ["Tool", "Where you use it", "What it does"],
        [
            ["Wedding Cost Prediction", "Budget → Plan your wedding budget", "Estimates vendor cost from your form (style, guests, district, ceremony, packages). Shows a range and category breakdown. If the estimate service is down, use Try again or continue with the on-page engine."],
            ["Vendor Match", "Vendors → Match", "Ranks listings by district, budget overlap and ceremony type. Not a machine-learning model — a match score. Needs district, budget and ceremony on your profile."],
            ["Smart Seating", "Seating Chart → Auto-seat all", "Places Coming guests using groups and “do not sit with” names. Review after Auto-seat flags are guidance. Falls back to rule-based seating if needed."],
            ["Readiness Score", "Home and Analytics", "Checklist 40% + RSVPs 35% + Budget 25%. On track at 70% or above."],
        ],
        col_widths=[4.2, 4.3, 8],
    )
    callout(
        doc,
        "Note",
        "Cost prediction and smart seating work best when the planning services are running. The rest of WowWed (checklist, guests, vendors, invitations, reports) works without them.",
        fill=NOTE_HEX,
        accent=GOLD_HEX,
    )

    # =====================================================================
    # 15 TIPS / TROUBLESHOOT / GLOSSARY
    # =====================================================================
    heading1(doc, "Tips, troubleshooting and glossary", "15")
    heading2(doc, "Recommended order for couples", "15.1")
    numbered(doc, "Create your account and save the wedding profile (names, date, district, ceremony, budget).")
    numbered(doc, "Skim the seeded Wedding Checklist and assign a few roles on Wedding Crew.")
    numbered(doc, "Build the Guest List (add many or import CSV) and mark Coming as replies arrive.")
    numbered(doc, "Run Plan your wedding budget, save a scenario, then log expenses.")
    numbered(doc, "Use Match on Vendors, send requests, and Confirm booking when you are ready.")
    numbered(doc, "When most Coming guests are in, open Seating Chart → Auto-seat all, then adjust.")
    numbered(doc, "Design invitations and download PDFs from Reports.")

    heading2(doc, "Good habits", "15.2")
    bullet(doc, "Keep the wedding profile current — it feeds almost every smart tool.")
    bullet(doc, "Only Coming guests are auto-seated; Waiting guests stay off the chart until you change RSVP.")
    bullet(doc, "Use Do not sit with on the guest card before you auto-seat.")
    bullet(doc, "Confirm booking is a planning confirmation, not a card charge.")
    bullet(doc, "Export CSVs and PDF reports before you make large list edits.")
    bullet(doc, "Vendors: finish photos and packages so Listing strength is high and couples trust the page.")

    heading2(doc, "If something looks wrong", "15.3")
    make_table(
        doc,
        ["What you see", "What to try"],
        [
            ["Cannot open the dashboard", "Confirm you used a couple account. Vendors are sent to /vendor."],
            ["Cannot open vendor pages", "Log in with the vendor account created via Join as a vendor."],
            ["Forgot password", "Login → Forgot password? → Update password (6+ characters)."],
            ["Empty checklist", "Save the wedding profile once so starter tasks can be created."],
            ["Auto-seat left people out", "Mark them Coming. Add tables or raise chairs per table."],
            ["Cost estimate error", "Click Try again, or use Reset to wedding profile and continue tracking expenses."],
            ["Match shows poor vendors", "Set district, Total budget (LKR) and ceremony type on the profile."],
            ["Vendor date rejected", "Pick another date — the calendar may already hold a booking."],
            ["PDF download is a print window", "In the print dialog choose Save as PDF (or Microsoft Print to PDF)."],
            ["WowBot seems generic", "Log in as a couple and ask about your wedding, guests or budget so it can use live numbers."],
            ["Photo or PDF upload fails (vendor)", "Keep photos at or under 900 KB (max 6). Package PDFs at or under 2 MB."],
        ],
        col_widths=[5.5, 11],
    )

    heading2(doc, "Glossary", "15.4")
    make_table(
        doc,
        ["Term", "Meaning on WowWed"],
        [
            ["Coming / Waiting / Not coming", "RSVP states on the Guest List."],
            ["Confirm booking", "Couple action that marks a hire Paid and adds it to Budget."],
            ["Listing strength", "How complete a vendor’s public page is."],
            ["Match", "Rule-based vendor ranking from your profile."],
            ["On track", "Readiness score of 70% or higher."],
            ["Paid (booking)", "Confirmed hire in the planner — not an online card payment."],
            ["Readiness score", "Checklist 40% + RSVPs 35% + Budget 25%."],
            ["Review after Auto-seat", "Optional seating suggestions after smart seating."],
            ["Scenario", "A named budget plan you can save and reload."],
            ["WowBot", "Offline keyword assistant for planning questions."],
        ],
        col_widths=[5.5, 11],
    )

    heading2(doc, "Need more help", "15.5")
    para(
        doc,
        "Use Chat with WowBot or the Assistant page for in-app guidance. For account issues, use Reset password on Settings or the login page. "
        "Keep this manual with your planning notes so button names stay easy to find.",
        align="justify",
    )

    # Closing band
    para(doc, "", space_after=10)
    bar = doc.add_table(rows=1, cols=1)
    set_table_width(bar, 16.5)
    c = bar.cell(0, 0)
    shade_cell(c, TERRACOTTA_HEX)
    set_cell_borders(c, TERRACOTTA_HEX, "0")
    set_cell_margins(c, 100, 100, 140, 140)
    c.paragraphs[0].clear()
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.paragraphs[0].add_run("WowWed  ·  Plan Your Dream Wedding, Smarter.")
    set_run_font(r, size=12, bold=True, color=RGBColor(255, 255, 255))
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("User Manual  ·  Version 1.0  ·  " + date.today().strftime("%B %Y"))
    set_run_font(r, size=10, color=RGBColor(255, 240, 232))

    out = Path(__file__).resolve().parent / "WowWed_User_Manual.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
